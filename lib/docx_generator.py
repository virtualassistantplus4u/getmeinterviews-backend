import io
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


FONT_NAME = "Calibri"


def _add_border_bottom(paragraph):
    """Add a bottom border to a paragraph (for section headings)."""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "aaaaaa")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_spacing(paragraph, before_pt=0, after_pt=0):
    """Set paragraph spacing in points."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before_pt)
    pf.space_after = Pt(after_pt)
    pf.line_spacing = Pt(13)


def _add_text_runs(paragraph, text: str, base_size: float, force_bold: bool = False):
    """
    Parse **bold** markers and add styled runs.
    Strips em dashes as a safety net.
    """
    if not text:
        return
    # Safety: replace em dashes with commas
    text = text.replace("—", ",").replace("\u2014", ",")
    # Split on **bold** markers
    parts = re.split(r"\*\*([^*]+)\*\*", text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        run.font.name = FONT_NAME
        run.font.size = Pt(base_size)
        run.bold = force_bold or (i % 2 == 1)


def _section_heading(doc: Document, title: str):
    """Add an uppercase section heading with bottom border."""
    para = doc.add_paragraph()
    _set_spacing(para, before_pt=8, after_pt=3)
    _add_border_bottom(para)
    run = para.add_run(title.upper())
    run.font.name = FONT_NAME
    run.font.size = Pt(10)
    run.bold = True
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return para


def generate_docx(resume_data: dict) -> bytes:
    """
    Generate a fully formatted ATS-ready DOCX from resume JSON.
    Returns bytes ready to send as a file download.
    """
    doc = Document()

    # ── Page margins ──────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    # ── Default style ─────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(11)

    # ── Name ─────────────────────────────────────────────────
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(name_para, before_pt=0, after_pt=2)
    name_run = name_para.add_run(resume_data.get("name", "Candidate"))
    name_run.font.name = FONT_NAME
    name_run.font.size = Pt(18)
    name_run.bold = True
    name_run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)

    # ── Contact ───────────────────────────────────────────────
    if resume_data.get("contact"):
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_spacing(contact_para, before_pt=0, after_pt=8)
        contact_run = contact_para.add_run(resume_data["contact"])
        contact_run.font.name = FONT_NAME
        contact_run.font.size = Pt(9)
        contact_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # ── Summary ───────────────────────────────────────────────
    if resume_data.get("summary"):
        _section_heading(doc, "Professional Summary")
        para = doc.add_paragraph()
        _set_spacing(para, before_pt=3, after_pt=3)
        _add_text_runs(para, resume_data["summary"], 11)

    # ── Skills ────────────────────────────────────────────────
    skills = resume_data.get("skills", [])
    if skills:
        _section_heading(doc, "Core Competencies")
        para = doc.add_paragraph()
        _set_spacing(para, before_pt=3, after_pt=3)
        for i, skill in enumerate(skills):
            _add_text_runs(para, skill, 11)
            if i < len(skills) - 1:
                sep_run = para.add_run("   \u2022   ")
                sep_run.font.name = FONT_NAME
                sep_run.font.size = Pt(11)
                sep_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # ── Experience ────────────────────────────────────────────
    experience = resume_data.get("experience", [])
    if experience:
        _section_heading(doc, "Experience")
        for ei, exp in enumerate(experience):
            # Role + Company + Years on one line
            exp_para = doc.add_paragraph()
            _set_spacing(exp_para, before_pt=3 if ei == 0 else 8, after_pt=2)

            # Role (bold)
            role_run = exp_para.add_run(exp.get("role", ""))
            role_run.font.name = FONT_NAME
            role_run.font.size = Pt(11.5)
            role_run.bold = True

            # Company
            if exp.get("company"):
                company_run = exp_para.add_run(f"   \u00b7   {exp['company']}")
                company_run.font.name = FONT_NAME
                company_run.font.size = Pt(11)
                company_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

            # Years (right-aligned via tab)
            if exp.get("years"):
                tab_run = exp_para.add_run("\t")
                tab_run.font.name = FONT_NAME
                years_run = exp_para.add_run(exp["years"])
                years_run.font.name = FONT_NAME
                years_run.font.size = Pt(10)
                years_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

            # Bullets (max 5)
            bullets = exp.get("bullets", [])[:5]
            for bullet_text in bullets:
                clean = re.sub(r"^[-\u2022\u00b7]\s*", "", bullet_text)
                bullet_para = doc.add_paragraph(style="List Bullet")
                _set_spacing(bullet_para, before_pt=1, after_pt=1)
                _add_text_runs(bullet_para, clean, 11)

    # ── Education ─────────────────────────────────────────────
    education = resume_data.get("education", [])
    if education:
        _section_heading(doc, "Education")
        for edu in education:
            edu_para = doc.add_paragraph()
            _set_spacing(edu_para, before_pt=3, after_pt=2)

            degree_run = edu_para.add_run(edu.get("degree", ""))
            degree_run.font.name = FONT_NAME
            degree_run.font.size = Pt(11)
            degree_run.bold = True

            if edu.get("institution"):
                inst_run = edu_para.add_run(f"   \u00b7   {edu['institution']}")
                inst_run.font.name = FONT_NAME
                inst_run.font.size = Pt(11)
                inst_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

            if edu.get("year"):
                tab_run = edu_para.add_run("\t")
                tab_run.font.name = FONT_NAME
                year_run = edu_para.add_run(edu["year"])
                year_run.font.name = FONT_NAME
                year_run.font.size = Pt(10)
                year_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # ── Save to bytes ─────────────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
