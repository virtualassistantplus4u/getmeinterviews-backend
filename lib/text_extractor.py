import pdfplumber
import docx
import io
from fastapi import HTTPException


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF bytes using pdfplumber."""
    try:
        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        result = "\n\n".join(text_parts).strip()
        if not result:
            raise HTTPException(status_code=422, detail="Could not extract text from PDF. Please ensure it is not a scanned image.")
        return result
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"PDF extraction failed: {str(e)}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text.strip())
        result = "\n".join(text_parts).strip()
        if not result:
            raise HTTPException(status_code=422, detail="Could not extract text from DOCX. The file may be empty.")
        return result
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"DOCX extraction failed: {str(e)}")


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Auto-detect file type and extract text."""
    filename_lower = filename.lower()
    if filename_lower.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif filename_lower.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    else:
        raise HTTPException(
            status_code=422,
            detail="Unsupported file type. Please upload a PDF or DOCX file."
        )
